import io
from datetime import datetime, timezone
from xml.sax.saxutils import escape as _xml_escape, quoteattr as _xml_quoteattr

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
)

from docx import Document
from docx.shared import Pt, RGBColor

from backend.models.schemas import InteractionResult

RISK_HEX = {"high": "#991b1b", "moderate": "#92400e", "low": "#065f46", "unknown": "#475569"}
RISK_RGB = {
    "high": RGBColor(0x99, 0x1B, 0x1B), "moderate": RGBColor(0x92, 0x40, 0x0E),
    "low": RGBColor(0x06, 0x5F, 0x46), "unknown": RGBColor(0x47, 0x55, 0x69),
}


def _timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _esc(text: str) -> str:
    """Escape text for ReportLab's mini-XML markup — every dynamic string
    (drug names, URLs, LLM-authored text) must go through this before being
    interpolated into a Paragraph. A source URL containing a literal quote
    character (seen in practice: openfda.generic_name:"warfarin") broke the
    attribute parser before this was added — not a hypothetical edge case."""
    return _xml_escape(text)


def _esc_attr(text: str) -> str:
    """Escape + quote an attribute value (e.g. an href) for ReportLab markup."""
    return _xml_quoteattr(text)


def _data_vintage_lines(result: InteractionResult) -> list[str]:
    """Plain-text lines noting when each cached data source was retrieved —
    matters if a user cites a result later. CLAUDE.md "Next Up" #3."""
    lines = []
    for drug in (result.drug_a, result.drug_b):
        if drug.resolved_at:
            lines.append(f"{drug.standard_name} — RxCUI resolved (RxNorm): {drug.resolved_at} UTC")
        if drug.structure_retrieved_at:
            lines.append(f"{drug.standard_name} — structure retrieved (PubChem): {drug.structure_retrieved_at} UTC")
        if drug.verified_mechanisms and drug.verified_mechanisms[0].retrieved_at:
            lines.append(
                f"{drug.standard_name} — mechanism retrieved (ChEMBL): "
                f"{drug.verified_mechanisms[0].retrieved_at} UTC"
            )
    if result.verified_severity:
        lines.append(f"DDInter severity dataset version: {result.verified_severity.dataset_date}")
    return lines


def _describe_patient_context(pc) -> str:
    parts = []
    if pc.age is not None:
        parts.append(f"age {pc.age}")
    if pc.renal_function:
        parts.append(f"renal: {pc.renal_function}")
    if pc.hepatic_function:
        parts.append(f"hepatic: {pc.hepatic_function}")
    if pc.pregnant is not None:
        parts.append("pregnant" if pc.pregnant else "not pregnant")
    if pc.other_conditions:
        parts.append(pc.other_conditions)
    return ", ".join(parts)


def build_pdf(result: InteractionResult) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, spaceAfter=4)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=15)
    small = ParagraphStyle("Small", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor("#64748b"))
    headline_level = result.severity_comparison.display_level.value if result.severity_comparison else result.risk_level.value
    risk_style = ParagraphStyle(
        "Risk", parent=styles["Normal"], fontSize=12,
        textColor=colors.HexColor(RISK_HEX.get(headline_level, "#475569")), spaceAfter=10,
    )

    story = [
        Paragraph("Drug-Drug Interaction Report", h1),
        Paragraph(f"{_esc(result.drug_a.standard_name)} + {_esc(result.drug_b.standard_name)}", body),
        Paragraph(f"Risk level: {_esc(headline_level.upper())}", risk_style),
    ]

    if result.verified_severity:
        story.append(Paragraph(
            f"Verified severity ({_esc(result.verified_severity.source)}): "
            f"{_esc(result.verified_severity.level.upper())}   ·   "
            f"AI risk level: {_esc(result.risk_level.value.upper())}", body
        ))
        if result.severity_comparison:
            story.append(Paragraph(_esc(result.severity_comparison.explanation), small))
    else:
        story.append(Paragraph(
            "No DDInter-verified severity found for this pair — risk level above is AI-assessed only.", small
        ))
    if result.patient_context_used:
        story.append(Paragraph(
            f"Assessed with patient context: {_esc(_describe_patient_context(result.patient_context_used))}", small
        ))

    story += [
        Paragraph("Summary", h2),
        Paragraph(_esc(result.llm_summary), body),
        Paragraph("Mechanism — Verified (ChEMBL)", h2),
    ]

    for drug in (result.drug_a, result.drug_b):
        if drug.verified_mechanisms:
            for m in drug.verified_mechanisms:
                refs = ", ".join(
                    f'<a href={_esc_attr(r.ref_url)} color="blue">{_esc(r.ref_type)}</a>' for r in m.references
                ) or "no references"
                action = f" ({_esc(m.action_type)})" if m.action_type else ""
                story.append(Paragraph(
                    f"<b>{_esc(drug.standard_name)}</b> — target: {_esc(m.target)}{action}"
                    f"<br/>{_esc(m.mechanism_of_action)}<br/>References: {refs}",
                    body,
                ))
        else:
            story.append(Paragraph(
                f"<b>{_esc(drug.standard_name)}</b> — no independently verified mechanism found in ChEMBL.", body
            ))
        story.append(Spacer(1, 6))

    story.append(Paragraph("Mechanism — AI-synthesized (not independently verified)", h2))
    story.append(Paragraph(_esc(result.mechanism), body))
    targets = ", ".join(result.targets_involved) or "none"
    story.append(Paragraph(
        f"Type: {_esc(result.mechanism_type.value)} · Targets mentioned: {_esc(targets)}", body
    ))
    if result.pathway:
        story.append(Paragraph(_esc(result.pathway).replace("\n", "<br/>"), body))

    story.append(Paragraph("Clinical Effect", h2))
    story.append(Paragraph(_esc(result.clinical_effect), body))
    story.append(Paragraph("Recommendation", h2))
    story.append(Paragraph(_esc(result.recommendation), body))

    if result.action_convention:
        story.append(Paragraph("Suggested Action — General Convention", h2))
        story.append(Paragraph(
            f"<b>{_esc(result.action_convention.action)}</b> — {_esc(result.action_convention.description)}", body
        ))
        story.append(Paragraph(_esc(result.action_convention.basis), small))

    story.append(Paragraph("Sources", h2))
    if result.sources or result.verified_severity:
        items = []
        if result.verified_severity:
            items.append(ListItem(Paragraph(
                'DDInter 2.0 (severity rating, CC BY-NC-SA 4.0) — '
                '<a href="https://ddinter.scbdd.com" color="blue">https://ddinter.scbdd.com</a> — '
                f'dataset version {_esc(result.verified_severity.dataset_date)}', body
            )))
        items += [
            ListItem(Paragraph(f'{_esc(s.name)} — <a href={_esc_attr(s.url)} color="blue">{_esc(s.url)}</a>', body))
            for s in result.sources
        ]
        story.append(ListFlowable(items, bulletType="bullet"))
    else:
        story.append(Paragraph(
            "No independently-cited sources — relied on the LLM's own pharmacology knowledge.", body
        ))

    vintage_lines = _data_vintage_lines(result)
    if vintage_lines:
        story.append(Paragraph("Data Vintage", h2))
        story.append(ListFlowable(
            [ListItem(Paragraph(_esc(line), small)) for line in vintage_lines], bulletType="bullet"
        ))

    story.append(Spacer(1, 14))
    story.append(Paragraph(_esc(result.disclaimer), small))
    story.append(Paragraph(f"Generated by DDI Checker on {_timestamp()} — for research purposes only.", small))

    doc.build(story)
    return buf.getvalue()


def build_docx(result: InteractionResult) -> bytes:
    doc = Document()

    doc.add_heading("Drug-Drug Interaction Report", level=1)
    doc.add_paragraph(f"{result.drug_a.standard_name} + {result.drug_b.standard_name}")

    headline_level = result.severity_comparison.display_level.value if result.severity_comparison else result.risk_level.value
    risk_p = doc.add_paragraph()
    risk_run = risk_p.add_run(f"Risk level: {headline_level.upper()}")
    risk_run.bold = True
    risk_run.font.color.rgb = RISK_RGB.get(headline_level, RGBColor(0x47, 0x55, 0x69))

    if result.verified_severity:
        doc.add_paragraph(
            f"Verified severity ({result.verified_severity.source}): {result.verified_severity.level.upper()}"
            f"   ·   AI risk level: {result.risk_level.value.upper()}"
        )
        if result.severity_comparison:
            note = doc.add_paragraph().add_run(result.severity_comparison.explanation)
            note.italic = True
            note.font.size = Pt(9)
    else:
        note = doc.add_paragraph().add_run(
            "No DDInter-verified severity found for this pair — risk level above is AI-assessed only."
        )
        note.italic = True
        note.font.size = Pt(9)
    if result.patient_context_used:
        ctx = doc.add_paragraph().add_run(
            f"Assessed with patient context: {_describe_patient_context(result.patient_context_used)}"
        )
        ctx.italic = True
        ctx.font.size = Pt(9)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(result.llm_summary)

    doc.add_heading("Mechanism — Verified (ChEMBL)", level=2)
    for drug in (result.drug_a, result.drug_b):
        doc.add_paragraph().add_run(drug.standard_name).bold = True
        if drug.verified_mechanisms:
            for m in drug.verified_mechanisms:
                label = f"Target: {m.target}" + (f" ({m.action_type})" if m.action_type else "")
                doc.add_paragraph(label, style="List Bullet")
                doc.add_paragraph(m.mechanism_of_action, style="List Bullet")
                for r in m.references:
                    doc.add_paragraph(f"{r.ref_type}: {r.ref_url}", style="List Bullet")
        else:
            doc.add_paragraph("No independently verified mechanism found in ChEMBL.", style="List Bullet")

    doc.add_heading("Mechanism — AI-synthesized (not independently verified)", level=2)
    doc.add_paragraph(result.mechanism)
    doc.add_paragraph(
        f"Type: {result.mechanism_type.value} · Targets mentioned: {', '.join(result.targets_involved) or 'none'}"
    )
    if result.pathway:
        doc.add_paragraph(result.pathway)

    doc.add_heading("Clinical Effect", level=2)
    doc.add_paragraph(result.clinical_effect)
    doc.add_heading("Recommendation", level=2)
    doc.add_paragraph(result.recommendation)

    if result.action_convention:
        doc.add_heading("Suggested Action — General Convention", level=2)
        action_p = doc.add_paragraph()
        action_p.add_run(result.action_convention.action).bold = True
        action_p.add_run(f" — {result.action_convention.description}")
        basis_run = doc.add_paragraph().add_run(result.action_convention.basis)
        basis_run.italic = True
        basis_run.font.size = Pt(9)

    doc.add_heading("Sources", level=2)
    if result.sources or result.verified_severity:
        if result.verified_severity:
            doc.add_paragraph(
                "DDInter 2.0 (severity rating, CC BY-NC-SA 4.0) — https://ddinter.scbdd.com — "
                f"dataset version {result.verified_severity.dataset_date}",
                style="List Bullet"
            )
        for s in result.sources:
            doc.add_paragraph(f"{s.name} — {s.url}", style="List Bullet")
    else:
        doc.add_paragraph("No independently-cited sources — relied on the LLM's own pharmacology knowledge.")

    vintage_lines = _data_vintage_lines(result)
    if vintage_lines:
        doc.add_heading("Data Vintage", level=2)
        for line in vintage_lines:
            p = doc.add_paragraph(line, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()
    disclaimer_run = doc.add_paragraph().add_run(result.disclaimer)
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(9)

    footer_run = doc.add_paragraph().add_run(
        f"Generated by DDI Checker on {_timestamp()} — for research purposes only."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
